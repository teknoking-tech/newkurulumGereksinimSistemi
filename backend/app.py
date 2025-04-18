from flask import Flask, request, jsonify, send_file, render_template, make_response
from flask_cors import CORS
import os
import json
import logging
import requests
from datetime import datetime
import tempfile
import markdown
from docx import Document
import io
from weasyprint import HTML
import uuid
import jwt
from functools import wraps

# Loglama yapılandırması
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)  # Cross-Origin Resource Sharing etkinleştir

# Uygulama konfigürasyonu
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'cbot-secret-key-change-in-production')
app.config['JWT_EXPIRATION'] = 3600  # Token süresi (saniye)

# n8n webhook URL'i (gerçek uygulamada env dosyasından alınabilir)
N8N_WEBHOOK_URL = os.environ.get('N8N_WEBHOOK_URL', 'http://n8n:5678/webhook/cbot-chatbot')

# Mock veritabanı (gerçek uygulamada gerçek bir veritabanı kullanılır)
DATABASE = {
    'users': {
        'user@example.com': {'password': 'password123', 'role': 'user'},
        'admin@example.com': {'password': 'admin123', 'role': 'admin'}
    },
    'requirements': []
}

# Modül bilgileri
MODULE_INFO = {
    'core': {
        'name': 'Core',
        'description': 'Ana CBOT çekirdek modülü',
        'cpu_test': 4,
        'cpu_live': 8,
        'ram_test': 8,
        'ram_live': 16,
        'disk_test': 20,
        'disk_live': 40,
        'port': 5351
    },
    'panel': {
        'name': 'Panel',
        'description': 'Yönetim paneli arayüzü',
        'cpu_test': 2,
        'cpu_live': 4,
        'ram_test': 4,
        'ram_live': 8,
        'disk_test': 10,
        'disk_live': 20,
        'port': 3000
    },
    'livechat': {
        'name': 'Live Chat',
        'description': 'Canlı destek modülü',
        'cpu_test': 2,
        'cpu_live': 4,
        'ram_test': 4,
        'ram_live': 8,
        'disk_test': 10,
        'disk_live': 20,
        'port': 3200
    },
    'fusion': {
        'name': 'Fusion',
        'description': 'Entegrasyon modülü',
        'cpu_test': 2,
        'cpu_live': 6,
        'ram_test': 6,
        'ram_live': 12,
        'disk_test': 15,
        'disk_live': 30,
        'port': 9600
    },
    'classifier': {
        'name': 'Classifier',
        'description': 'Sınıflandırma modülü (MS SQL Server gerektirir)',
        'cpu_test': 8,
        'cpu_live': 16,
        'ram_test': 16,
        'ram_live': 32,
        'disk_test': 30,
        'disk_live': 60,
        'port': 9100,
        'db_requirement': 'mssql'
    },
    'aiflow': {
        'name': 'AI Flow',
        'description': 'Yapay zeka akış modülü',
        'cpu_test': 10,
        'cpu_live': 20,
        'ram_test': 16,
        'ram_live': 32,
        'disk_test': 40,
        'disk_live': 80,
        'port': 8500
    },
    'analytics': {
        'name': 'Analytics',
        'description': 'Analitik modülü',
        'cpu_test': 4,
        'cpu_live': 8,
        'ram_test': 8,
        'ram_live': 16,
        'disk_test': 20,
        'disk_live': 40,
        'port': 7500
    }
}

# Yardımcı servis bilgileri
SERVICE_INFO = {
    'ocr': {
        'name': 'OCR',
        'description': 'Optik Karakter Tanıma servisi',
        'cpu_test': 4,
        'cpu_live': 8,
        'ram_test': 8,
        'ram_live': 16,
        'disk_test': 10,
        'disk_live': 20,
        'requires_gpu': False
    },
    'masking': {
        'name': 'Maskeleme',
        'description': 'Veri maskeleme servisi',
        'cpu_test': 2,
        'cpu_live': 4,
        'ram_test': 4,
        'ram_live': 8,
        'disk_test': 5,
        'disk_live': 10,
        'requires_gpu': False
    },
    'file_to_md': {
        'name': 'File to Markdown',
        'description': 'Dosya içeriğini Markdown formatına çevirme servisi',
        'cpu_test': 2,
        'cpu_live': 4,
        'ram_test': 4,
        'ram_live': 8,
        'disk_test': 5,
        'disk_live': 10,
        'requires_gpu': False
    },
    'crawler': {
        'name': 'Crawler',
        'description': 'Web sitesi tarama servisi',
        'cpu_test': 4,
        'cpu_live': 8,
        'ram_test': 8,
        'ram_live': 16,
        'disk_test': 15,
        'disk_live': 30,
        'requires_gpu': False
    },
    'hf_model_hosting': {
        'name': 'HF Model Hosting',
        'description': 'Huggingface model barındırma servisi',
        'cpu_test': 8,
        'cpu_live': 16,
        'ram_test': 16,
        'ram_live': 32,
        'disk_test': 30,
        'disk_live': 60,
        'requires_gpu': True,
        'gpu_ram': 16
    }
}

# Veritabanı bilgileri
DB_INFO = {
    'mongodb': {
        'name': 'MongoDB',
        'port': 27017,
        'collation': None
    },
    'mssql': {
        'name': 'Microsoft SQL Server',
        'port': 1433,
        'collation': 'SQL_Latin1_General_CP1_CI_AS'
    },
    'postgresql': {
        'name': 'PostgreSQL',
        'port': 5432,
        'collation': None
    }
}

# JWT token doğrulama decoratörü
def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        
        # Token'ı header'dan al
        if 'Authorization' in request.headers:
            auth_header = request.headers['Authorization']
            if auth_header.startswith('Bearer '):
                token = auth_header.split(' ')[1]
        
        if not token:
            return jsonify({'message': 'Token bulunamadı!'}), 401
        
        try:
            # Token'ı doğrula
            data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
            current_user = {'email': data['email'], 'role': data['role']}
        except:
            return jsonify({'message': 'Token geçersiz!'}), 401
        
        return f(current_user, *args, **kwargs)
    
    return decorated

# Admin yetkisi kontrol decoratörü
def admin_required(f):
    @wraps(f)
    def decorated(current_user, *args, **kwargs):
        if current_user['role'] != 'admin':
            return jsonify({'message': 'Bu işlem için admin yetkisi gereklidir!'}), 403
        return f(current_user, *args, **kwargs)
    
    return decorated

# Ana sayfa
@app.route('/')
def index():
    return app.send_static_file('index.html')

# Kullanıcı girişi
@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')
    role = data.get('role', 'user')  # Varsayılan olarak user rolü
    
    if not email or not password:
        return jsonify({'message': 'Email ve şifre gereklidir!'}), 400
    
    # Kullanıcı kimlik doğrulama
    if email not in DATABASE['users'] or DATABASE['users'][email]['password'] != password:
        return jsonify({'message': 'Geçersiz kullanıcı bilgileri!'}), 401
    
    # Kullanıcı rolü kontrolü
    if role == 'admin' and DATABASE['users'][email]['role'] != 'admin':
        return jsonify({'message': 'Admin yetkisi bulunamadı!'}), 403
    
    # JWT token oluştur
    token = jwt.encode(
        {
            'email': email,
            'role': DATABASE['users'][email]['role'],
            'exp': datetime.utcnow().timestamp() + app.config['JWT_EXPIRATION']
        },
        app.config['SECRET_KEY'],
        algorithm='HS256'
    )
    
    return jsonify({
        'message': 'Giriş başarılı!',
        'token': token,
        'user': {
            'email': email,
            'role': DATABASE['users'][email]['role']
        }
    }), 200

# Chatbot mesaj işleme
@app.route('/api/chatbot/message', methods=['POST'])
def process_message():
    data = request.get_json()
    user_message = data.get('message')
    
    if not user_message:
        return jsonify({'message': 'Mesaj içeriği gereklidir!'}), 400
    
    try:
        # n8n webhook'a mesajı ilet
        webhook_response = requests.post(
            N8N_WEBHOOK_URL,
            json={'message': user_message},
            timeout=5
        )
        
        if webhook_response.status_code == 200:
            response_data = webhook_response.json()
            bot_message = response_data.get('response', 'Mesajınız alındı, ancak bir cevap oluşturulamadı.')
        else:
            # Webhook başarısız olursa basit bir cevap döndür
            logger.warning(f"n8n webhook failed with status {webhook_response.status_code}")
            bot_message = "Şu anda sisteme ulaşılamıyor. Lütfen daha sonra tekrar deneyin."
    
    except requests.RequestException as e:
        logger.error(f"Error connecting to n8n webhook: {e}")
        bot_message = "Chatbot servisine bağlanırken bir hata oluştu. Lütfen daha sonra tekrar deneyin."
    
    return jsonify({
        'message': bot_message
    }), 200

# Gereksinim dokümanı oluşturma
@app.route('/api/requirements/generate', methods=['POST'])
def generate_requirements():
    data = request.get_json()
    
    # Gerekli alanların kontrolü
    required_fields = ['environment', 'environmentType', 'coreModules', 'database']
    for field in required_fields:
        if field not in data:
            return jsonify({'message': f'{field} alanı gereklidir!'}), 400
    
    # Seçilen modül ve servisleri birleştir
    selected_modules = data.get('coreModules', [])
    selected_services = data.get('auxiliaryServices', [])
    
    # Veritabanı kontrolü
    database = data.get('database')
    if database not in DB_INFO:
        return jsonify({'message': 'Geçersiz veritabanı seçimi!'}), 400
    
    # Classifier modülü için MSSQL kontrolü
    if 'classifier' in selected_modules and database != 'mssql':
        return jsonify({
            'message': 'Classifier modülü yalnızca MSSQL veritabanı ile çalışmaktadır!',
            'warning': True
        }), 200  # İşlemi devam ettirmek için 200 döndürüyoruz, frontend tarafında uyarı gösterilecek
    
    # Gereksinim belgesi oluşturma
    requirements = generate_requirements_document(data)
    
    # Benzersiz bir ID ile gereksinimleri kaydet
    req_id = str(uuid.uuid4())
    requirements['id'] = req_id
    requirements['created_at'] = datetime.utcnow().isoformat()
    DATABASE['requirements'].append(requirements)
    
    return jsonify({
        'message': 'Gereksinim dokümanı başarıyla oluşturuldu.',
        'requirementId': req_id,
        'requirements': requirements
    }), 201

# PDF dokümanı indirme
@app.route('/api/requirements/<req_id>/pdf', methods=['GET'])
def download_pdf(req_id):
    # Gereksinim dokümanını bul
    requirement = next((r for r in DATABASE['requirements'] if r.get('id') == req_id), None)
    if not requirement:
        return jsonify({'message': 'Gereksinim dokümanı bulunamadı!'}), 404
    
    # HTML içeriğini hazırla
    html_content = render_template(
        'requirement_template.html',
        requirement=requirement
    )
    
    # PDF oluştur
    pdf_file = io.BytesIO()
    HTML(string=html_content).write_pdf(pdf_file)
    pdf_file.seek(0)
    
    # PDF'i indir
    response = make_response(pdf_file.getvalue())
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename=CBOT_Kurulum_Gereksinimleri_{req_id}.pdf'
    
    return response

# Word dokümanı indirme
@app.route('/api/requirements/<req_id>/docx', methods=['GET'])
def download_docx(req_id):
    # Gereksinim dokümanını bul
    requirement = next((r for r in DATABASE['requirements'] if r.get('id') == req_id), None)
    if not requirement:
        return jsonify({'message': 'Gereksinim dokümanı bulunamadı!'}), 404
    
    # Word belgesi oluştur
    doc = Document()
    doc.add_heading('CBOT Kurulum Gereksinimleri', 0)
    
    # Belge içeriğini doldur
    add_requirement_content_to_docx(doc, requirement)
    
    # Dosyayı kaydet
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    # Word belgesini indir
    response = make_response(docx_file.getvalue())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = f'attachment; filename=CBOT_Kurulum_Gereksinimleri_{req_id}.docx'
    
    return response

# Admin paneli için modül yapılandırması güncelleme
@app.route('/api/admin/modules/<module_id>', methods=['PUT'])
@token_required
@admin_required
def update_module(current_user, module_id):
    if module_id not in MODULE_INFO:
        return jsonify({'message': 'Modül bulunamadı!'}), 404
    
    data = request.get_json()
    
    # Güncellenebilir alanlar
    updatable_fields = ['cpu_test', 'cpu_live', 'ram_test', 'ram_live', 'disk_test', 'disk_live', 'description']
    
    # Modül bilgilerini güncelle
    for field in updatable_fields:
        if field in data:
            MODULE_INFO[module_id][field] = data[field]
    
    return jsonify({
        'message': 'Modül başarıyla güncellendi.',
        'module': MODULE_INFO[module_id]
    }), 200

# Gereksinim dokümanı oluşturma yardımcı fonksiyonu
def generate_requirements_document(data):
    environment = data.get('environment')
    environment_type = data.get('environmentType')
    core_modules = data.get('coreModules', [])
    auxiliary_services = data.get('auxiliaryServices', [])
    database = data.get('database')
    ldap_enabled = data.get('ldapEnabled', False)
    ldap_details = data.get('ldapDetails', {})
    
    # Donanım gereksinimlerini hesapla
    hardware_requirements = calculate_hardware_requirements(
        environment, core_modules, auxiliary_services
    )
    
    # Veritabanı gereksinimlerini hazırla
    database_requirements = {
        'type': DB_INFO[database]['name'],
        'port': DB_INFO[database]['port'],
        'collation': DB_INFO[database]['collation']
    }
    
    # Network/Firewall gereksinimlerini hazırla
    network_requirements = {
        'internet_access': True,
        'db_port': DB_INFO[database]['port'],
        'dns_records': generate_dns_records(environment, core_modules)
    }
    
    # Docker gereksinimleri
    docker_requirements = {
        'registry': 'registry.cbot.ai',
        'port_access': 443,
        'containers': []
    }
    
    # Container listesini oluştur
    for module in core_modules:
        docker_requirements['containers'].append(f'cbot-{module}')
    
    for service in auxiliary_services:
        docker_requirements['containers'].append(f'cbot-{service.replace("_", "-")}')
    
    # Gereksinim belgesini oluştur
    requirements_doc = {
        'environment': environment,
        'environment_type': environment_type,
        'core_modules': [{'id': m, 'name': MODULE_INFO[m]['name']} for m in core_modules],
        'auxiliary_services': [{'id': s, 'name': SERVICE_INFO[s]['name']} for s in auxiliary_services],
        'hardware_requirements': hardware_requirements,
        'database_requirements': database_requirements,
        'network_requirements': network_requirements,
        'docker_requirements': docker_requirements,
        'created_at': datetime.utcnow().isoformat()
    }
    
    # LDAP gereksinimleri ekle
    if ldap_enabled:
        requirements_doc['ldap_requirements'] = {
            'url': ldap_details.get('url', ''),
            'bind_dn': ldap_details.get('bindDn', ''),
            'search_base': ldap_details.get('searchBase', ''),
            'search_filter': ldap_details.get('searchFilter', '')
        }
    
    # AI Flow veya HF Model Hosting için ek gereksinimler
    if 'aiflow' in core_modules or 'hf_model_hosting' in auxiliary_services:
        requirements_doc['ai_requirements'] = []
        
        if 'aiflow' in core_modules:
            requirements_doc['ai_requirements'].append({
                'service': 'AI Flow',
                'cpu_min': MODULE_INFO['aiflow']['cpu_test'],
                'ram_min': MODULE_INFO['aiflow']['ram_test'],
                'disk_min': MODULE_INFO['aiflow']['disk_test'],
                'postgresql_disk': 4 if database == 'postgresql' else 0
            })
        
        if 'hf_model_hosting' in auxiliary_services:
            requirements_doc['ai_requirements'].append({
                'service': 'HF Model Hosting',
                'requires_gpu': True,
                'gpu_ram': SERVICE_INFO['hf_model_hosting']['gpu_ram'],
                'ram_min': SERVICE_INFO['hf_model_hosting']['ram_test']
            })
    
    return requirements_doc

# Donanım gereksinimlerini hesaplama
def calculate_hardware_requirements(environment, core_modules, auxiliary_services):
    # Başlangıç değerleri
    test_requirements = {
        'cpu': 4,  # Min CPU core
        'ram': 8,  # Min RAM (GB)
        'disk': 20,  # Min disk alanı (GB)
        'gpu': False,  # GPU gereksinimi
        'gpu_ram': 0  # GPU RAM gereksinimi (GB)
    }
    
    live_requirements = {
        'cpu': 8,  # Min CPU core
        'ram': 16,  # Min RAM (GB)
        'disk': 40,  # Min disk alanı (GB)
        'gpu': False,  # GPU gereksinimi
        'gpu_ram': 0  # GPU RAM gereksinimi (GB)
    }
    
    # Seçilen modüllere göre değerleri güncelle
    for module in core_modules:
        module_info = MODULE_INFO[module]
        test_requirements['cpu'] += module_info['cpu_test'] - 2  # Baz değerden düşerek ekle
        test_requirements['ram'] += module_info['ram_test'] - 4
        test_requirements['disk'] += module_info['disk_test'] - 10
        
        live_requirements['cpu'] += module_info['cpu_live'] - 4
        live_requirements['ram'] += module_info['ram_live'] - 8
        live_requirements['disk'] += module_info['disk_live'] - 20
    
    # Seçilen yardımcı servislere göre değerleri güncelle
    for service in auxiliary_services:
        service_info = SERVICE_INFO[service]
        test_requirements['cpu'] += service_info['cpu_test'] - 2
        test_requirements['ram'] += service_info['ram_test'] - 4
        test_requirements['disk'] += service_info['disk_test'] - 5
        
        live_requirements['cpu'] += service_info['cpu_live'] - 4
        live_requirements['ram'] += service_info['ram_live'] - 8
        live_requirements['disk'] += service_info['disk_live'] - 10
        
        # GPU gereksinimi kontrolü
        if service_info.get('requires_gpu', False):
            test_requirements['gpu'] = True
            live_requirements['gpu'] = True
            
            gpu_ram = service_info.get('gpu_ram', 0)
            test_requirements['gpu_ram'] = max(test_requirements['gpu_ram'], gpu_ram)
            live_requirements['gpu_ram'] = max(live_requirements['gpu_ram'], gpu_ram)
    
    # Minimum değerlerin altında olmadığından emin ol
    test_requirements['cpu'] = max(test_requirements['cpu'], 4)
    test_requirements['ram'] = max(test_requirements['ram'], 8)
    test_requirements['disk'] = max(test_requirements['disk'], 20)
    
    live_requirements['cpu'] = max(live_requirements['cpu'], 8)
    live_requirements['ram'] = max(live_requirements['ram'], 16)
    live_requirements['disk'] = max(live_requirements['disk'], 40)
    
    # Seçilen ortama göre sonuçları filtrele
    if environment == 'test':
        return {'test': test_requirements}
    elif environment == 'live':
        return {'live': live_requirements}
    else:  # Her ikisi
        return {'test': test_requirements, 'live': live_requirements}

# DNS kayıtları oluşturma
def generate_dns_records(environment, core_modules):
    dns_records = []
    
    # Modüllere göre DNS kayıtları oluştur
    dns_mapping = {
        'panel': {'name': 'cbot-panel', 'port': 3000},
        'core': {'name': 'cbot-core', 'port': 5351},
        'fusion': {'name': 'cbot-fusion', 'port': 9600},
        'livechat': {'name': 'cbot-livechat', 'port': 3200},
        'classifier': {'name': 'cbot-classifier', 'port': 9100},
        'aiflow': {'name': 'cbot-aiflow', 'port': 8500},
        'analytics': {'name': 'cbot-analytics', 'port': 7500}
    }
    
    # Core modülü varsa socket kaydı da ekle
    if 'core' in core_modules:
        dns_mapping['socket'] = {'name': 'cbot-socket', 'port': 5000, 'websocket': True}
    
    for module, info in dns_mapping.items():
        if module in core_modules or module == 'socket':
            if environment == 'test' or environment == 'both':
                dns_records.append({
                    'service': info['name'] + '-test',
                    'port': info['port'],
                    'websocket': info.get('websocket', False)
                })
            
            if environment == 'live' or environment == 'both':
                dns_records.append({
                    'service': info['name'],
                    'port': info['port'],
                    'websocket': info.get('websocket', False)
                })
    
    return dns_records

# Word belgesi içeriğini oluşturma
def add_requirement_content_to_docx(doc, requirement):
    # Belge başlığı ve tarihi
    doc.add_paragraph(f"Oluşturulma Tarihi: {datetime.fromisoformat(requirement['created_at']).strftime('%d.%m.%Y %H:%M')}")
    doc.add_heading('1. Ortam Bilgileri', level=1)
    
    # Ortam bilgileri
    environment_text = ''
    if requirement['environment'] == 'both':
        environment_text = 'Test ve Canlı'
    elif requirement['environment'] == 'test':
        environment_text = 'Test'
    else:
        environment_text = 'Canlı'
    
    environment_type_text = 'On-Premise' if requirement['environment_type'] == 'on-prem' else 'Bulut'
    
    p = doc.add_paragraph()
    p.add_run('Ortam Tipi: ').bold = True
    p.add_run(environment_text)
    
    p = doc.add_paragraph()
    p.add_run('Kurulum Ortamı: ').bold = True
    p.add_run(environment_type_text)
    
    # Seçilen modüller
    doc.add_heading('2. Seçilen Modüller', level=1)
    for module in requirement['core_modules']:
        doc.add_paragraph(module['name'], style='List Bullet')
    
    # Yardımcı servisler
    if requirement['auxiliary_services']:
        doc.add_heading('Yardımcı Servisler:', level=2)
        for service in requirement['auxiliary_services']:
            doc.add_paragraph(service['name'], style='List Bullet')
    
    # Donanım gereksinimleri
    doc.add_heading('3. Donanım Gereksinimleri', level=1)
    
    hardware_table = doc.add_table(rows=1, cols=6)
    hardware_table.style = 'Table Grid'
    
    # Tablo başlıkları
    header_cells = hardware_table.rows[0].cells
    header_cells[0].text = 'Ortam'
    header_cells[1].text = 'CPU'
    header_cells[2].text = 'RAM'
    header_cells[3].text = 'Disk'
    header_cells[4].text = 'GPU'
    header_cells[5].text = 'OS'
    
    # Tablo içeriği
    hw_requirements = requirement['hardware_requirements']
    
    if 'test' in hw_requirements:
        row_cells = hardware_table.add_row().cells
        row_cells[0].text = 'Test'
        row_cells[1].text = f"{hw_requirements['test']['cpu']} Core"
        row_cells[2].text = f"{hw_requirements['test']['ram']} GB"
        row_cells[3].text = f"{hw_requirements['test']['disk']} GB"
        row_cells[4].text = "Gerekmiyor" if not hw_requirements['test']['gpu'] else f"{hw_requirements['test']['gpu_ram']} GB VRAM"
        row_cells[5].text = "CentOS 7/8 veya RHEL"
    
    if 'live' in hw_requirements:
        row_cells = hardware_table.add_row().cells
        row_cells[0].text = 'Canlı'
        row_cells[1].text = f"{hw_requirements['live']['cpu']} Core"
        row_cells[2].text = f"{hw_requirements['live']['ram']} GB"
        row_cells[3].text = f"{hw_requirements['live']['disk']} GB"
        row_cells[4].text = "Gerekmiyor" if not hw_requirements['live']['gpu'] else f"{hw_requirements['live']['gpu_ram']} GB VRAM"
        row_cells[5].text = "CentOS 7/8 veya RHEL"
    
    # Veritabanı gereksinimleri
    doc.add_heading('4. Veritabanı Gereksinimleri', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Veritabanı Tipi: ').bold = True
    p.add_run(requirement['database_requirements']['type'])
    
    p = doc.add_paragraph()
    p.add_run('Kullanıcı: ').bold = True
    p.add_run('OWNER yetkili kullanıcı')
    
    if requirement['database_requirements']['collation']:
        p = doc.add_paragraph()
        p.add_run('Collation: ').bold = True
        p.add_run(requirement['database_requirements']['collation'])
    
    p = doc.add_paragraph()
    p.add_run('Port: ').bold = True
    p.add_run(str(requirement['database_requirements']['port']))
    
    doc.add_paragraph('IP ve port paylaşımı zorunludur')
    
    # Network/Firewall gereksinimleri
    doc.add_heading('5. Network / Firewall Gereksinimleri', level=1)
    doc.add_paragraph('Sunucular internet erişimli olmalıdır', style='List Bullet')
    doc.add_paragraph(f"Veritabanı portları açık olmalıdır ({requirement['database_requirements']['port']})", style='List Bullet')
    doc.add_paragraph('Load Balancer 5000 portunda WebSocket desteği sağlamalıdır', style='List Bullet')
    
    doc.add_heading('DNS Kayıtları:', level=2)
    dns_table = doc.add_table(rows=1, cols=3)
    dns_table.style = 'Table Grid'
    
    # DNS tablosu başlıkları
    dns_header = dns_table.rows[0].cells
    dns_header[0].text = 'Servis'
    dns_header[1].text = 'DNS Kaydı'
    dns_header[2].text = 'Port'
    
    # DNS tablosu içeriği
    for record in requirement['network_requirements']['dns_records']:
        row_cells = dns_table.add_row().cells
        row_cells[0].text = record['service'].split('-')[1].capitalize() + (' (Test)' if record['service'].endswith('-test') else '')
        row_cells[1].text = record['service']
        row_cells[2].text = f"{record['port']}{' (WebSocket)' if record.get('websocket', False) else ''}"
    
    # Docker Image ve Registry
    doc.add_heading('6. Docker Image ve Registry', level=1)
    doc.add_paragraph('Tüm servis image\'ları registry.cbot.ai adresinden çekilir', style='List Bullet')
    doc.add_paragraph('Sunucular 443 portundan bu adrese erişebilmelidir', style='List Bullet')
    
    doc.add_heading('Gerekli Containerlar:', level=2)
    for container in requirement['docker_requirements']['containers']:
        doc.add_paragraph(container, style='List Bullet')
    
    # LDAP Gereksinimleri (varsa)
    section_count = 7
    if 'ldap_requirements' in requirement:
        doc.add_heading(f'{section_count}. LDAP Gereksinimleri', level=1)
        
        p = doc.add_paragraph()
        p.add_run('LDAP URL: ').bold = True
        p.add_run(requirement['ldap_requirements']['url'] or 'Belirtilmedi')
        
        p = doc.add_paragraph()
        p.add_run('BIND DN: ').bold = True
        p.add_run(requirement['ldap_requirements']['bind_dn'] or 'Belirtilmedi')
        
        p = doc.add_paragraph()
        p.add_run('SEARCH BASE: ').bold = True
        p.add_run(requirement['ldap_requirements']['search_base'] or 'Belirtilmedi')
        
        p = doc.add_paragraph()
        p.add_run('SEARCH FILTER: ').bold = True
        p.add_run(requirement['ldap_requirements']['search_filter'] or 'Belirtilmedi')
        
        doc.add_paragraph('Email, Role, Name field\'ları mapping yapılmalıdır.')
        doc.add_paragraph('Rollerin panelde LDAP ile birebir eşleşmesi gerekmektedir.')
        
        section_count += 1
    
    # AI Servisleri Ek Gereksinimleri (varsa)
    if 'ai_requirements' in requirement:
        doc.add_heading(f'{section_count}. AI Servisleri Ek Gereksinimleri', level=1)
        
        for ai_req in requirement['ai_requirements']:
            doc.add_heading(f"{ai_req['service']} Gereksinimleri:", level=2)
            
            if ai_req['service'] == 'AI Flow':
                doc.add_paragraph(f"Minimum {ai_req['cpu_min']} Core CPU, {ai_req['ram_min']} GB RAM", style='List Bullet')
                if ai_req.get('postgresql_disk', 0) > 0:
                    doc.add_paragraph(f"{ai_req['postgresql_disk']} GB disk alanı (PostgreSQL)", style='List Bullet')
            
            if ai_req['service'] == 'HF Model Hosting':
                doc.add_paragraph('GPU destekli sunucu', style='List Bullet')
                doc.add_paragraph(f"Minimum {ai_req['gpu_ram']} GB VRAM", style='List Bullet')
                doc.add_paragraph(f"Minimum {ai_req['ram_min']} GB sistem RAM", style='List Bullet')
        
        section_count += 1
    
    # Kurulum Topolojisi
    doc.add_heading(f'{section_count}. Kurulum Topolojisi', level=1)
    doc.add_paragraph('Kurulum topolojisi görsel temsili AI asistan tarafından oluşturulmuştur.')
    
    # Doküman sonunda yasal uyarı
    doc.add_paragraph()
    doc.add_paragraph('Bu doküman CBOT tarafından otomatik olarak oluşturulmuştur. Detaylı bilgi için teknik ekibimizle iletişime geçebilirsiniz.')

# Ana uygulamayı başlat
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('DEBUG', 'False').lower() == 'true'
    
    app.run(host='0.0.0.0', port=port, debug=debug)